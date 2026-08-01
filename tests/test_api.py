from io import BytesIO

from fastapi.testclient import TestClient

import app.main as main
from app.database import DictionaryStore
from app.documents import DocumentSessionStore


def test_text_api_dictionary_and_static_app(tmp_path, monkeypatch) -> None:
    monkeypatch.setattr(main, "dictionary_store", DictionaryStore(tmp_path / "db.sqlite3"))
    monkeypatch.setattr(main, "session_store", DocumentSessionStore(tmp_path / "sessions"))

    with TestClient(main.app) as client:
        assert client.get("/").status_code == 200
        health = client.get("/api/health")
        assert health.status_code == 200
        assert health.json()["local_only"] is True

        created = client.post(
            "/api/dictionary",
            json={"term": "極秘案件", "entity_type": "CUSTOM", "note": "test"},
        )
        assert created.status_code == 201

        analyzed = client.post(
            "/api/analyze",
            json={"text": "極秘案件の連絡先はtaro@example.jp", "entities": ["CUSTOM", "EMAIL_ADDRESS"]},
        )
        assert analyzed.status_code == 200
        findings = analyzed.json()["findings"]
        assert {item["entity_type"] for item in findings} == {"CUSTOM", "EMAIL_ADDRESS"}

        masked = client.post(
            "/api/mask",
            json={
                "text": analyzed.json()["text"],
                "findings": findings,
                "accepted_ids": [item["id"] for item in findings],
                "mask_character": "█",
            },
        )
        assert masked.status_code == 200
        assert "極秘案件" not in masked.json()["masked_text"]
        assert "taro@example.jp" not in masked.json()["masked_text"]

        labeled = client.post(
            "/api/mask",
            json={
                "text": analyzed.json()["text"],
                "findings": findings,
                "accepted_ids": [item["id"] for item in findings],
                "replacement_mode": "entity_label",
            },
        )
        assert labeled.status_code == 200
        assert labeled.json()["masked_text"] == "[ユーザー定義]の連絡先は[メールアドレス]"

        manual = client.post(
            "/api/findings/manual",
            json={
                "text": "秘密と秘密",
                "findings": [],
                "block_id": "text",
                "start": 0,
                "end": 2,
                "entity_type": "CUSTOM",
                "scope": "all",
                "save_to_dictionary": True,
            },
        )
        assert manual.status_code == 200
        assert manual.json()["added_count"] == 2
        assert manual.json()["dictionary_status"] == "created"
        assert {item["source"] for item in manual.json()["added_findings"]} == {
            "manual-selection"
        }

        repeated = client.post(
            "/api/findings/manual",
            json={
                "text": "秘密",
                "findings": [],
                "block_id": "text",
                "start": 0,
                "end": 2,
                "entity_type": "CUSTOM",
                "scope": "single",
                "save_to_dictionary": True,
            },
        )
        assert repeated.status_code == 200
        assert repeated.json()["dictionary_status"] == "already_exists"

        deleted = client.delete(f"/api/dictionary/{created.json()['id']}")
        assert deleted.status_code == 204
        assert deleted.content == b""


def test_document_manual_finding_is_saved_and_masked(tmp_path, monkeypatch) -> None:
    from docx import Document

    monkeypatch.setattr(main, "dictionary_store", DictionaryStore(tmp_path / "db.sqlite3"))
    monkeypatch.setattr(main, "session_store", DocumentSessionStore(tmp_path / "sessions"))
    document = Document()
    source_text = "公開情報 極秘符号ZXCV 公開情報"
    document.add_paragraph(source_text)
    payload = BytesIO()
    document.save(payload)

    with TestClient(main.app) as client:
        analyzed = client.post(
            "/api/documents/analyze",
            files={
                "file": (
                    "manual.docx",
                    payload.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"entities": "EMAIL_ADDRESS"},
        )
        assert analyzed.status_code == 200
        analysis = analyzed.json()
        assert analysis["findings"] == []
        block = next(item for item in analysis["blocks"] if item["text"] == source_text)
        start = source_text.index("極秘符号ZXCV")

        manual = client.post(
            f"/api/documents/{analysis['session_id']}/findings/manual",
            json={
                "block_id": block["id"],
                "start": start,
                "end": start + len("極秘符号ZXCV"),
                "entity_type": "CUSTOM",
                "scope": "single",
                "save_to_dictionary": False,
            },
        )
        assert manual.status_code == 200
        finding = manual.json()["added_findings"][0]
        _, _, saved_findings = main.session_store.load(analysis["session_id"])
        assert [item.id for item in saved_findings] == [finding["id"]]

        masked = client.post(
            f"/api/documents/{analysis['session_id']}/mask",
            json={"accepted_ids": [finding["id"]], "mask_character": "█"},
        )
        assert masked.status_code == 200
        downloaded = client.get(masked.json()["download_url"])
        masked_document = Document(BytesIO(downloaded.content))
        assert masked_document.paragraphs[0].text == "公開情報 ████████ 公開情報"

        removed = client.delete(
            f"/api/documents/{analysis['session_id']}/findings/{finding['id']}"
        )
        assert removed.status_code == 204
        _, _, remaining_findings = main.session_store.load(analysis["session_id"])
        assert remaining_findings == []

        missing = client.post(
            "/api/documents/00000000000000000000000000000000/findings/manual",
            json={
                "block_id": block["id"],
                "start": start,
                "end": start + 2,
                "entity_type": "CUSTOM",
            },
        )
        assert missing.status_code == 404
